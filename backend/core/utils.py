import random
from django.conf import settings
from django.core.mail import send_mail
from django.core.validators import validate_email
from django.core.exceptions import ValidationError

import os
import PyPDF2
import pdfplumber
from docx import Document


def generate_otp():
    return str(random.randint(100000, 999999))

def send_otp_via_email(email, otp):

    try:
        validate_email(email)
    except ValidationError:
        print("Invalid email format")
        return False
    
    subject = "Your OTP for Login/Registration"
    message = f"Your One-Time Password (OTP) is: {otp}\nThis OTP is valid for 5 minutes."
    # from_email = settings.DEFAULT_FROM_EMAIL  
    from_email = None
    
    try:
        send_mail(subject, message, from_email, [email])
        return True
    except Exception as e:
        print(f"Error sending email: {e}")
        return False
    

# Helper functions for text extraction
def extract_text_from_pdf(file_path):
    text = ""
    urls = []  # List to collect URLs
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # Extract text
            text += page.extract_text() or ""

            # Extract hyperlinks (annotations)
            if page.annots:
                for annot in page.annots:
                    if annot.get("uri"):  # Check if a URL exists
                        urls.append(annot.get("uri"))
    
    return text.strip(), urls


# def extract_text_from_word(file_path):
#     document = Document(file_path)
#     text = []
#     urls = []  # List to collect URLs

#     # Loop through paragraphs to get text and hyperlinks
#     for para in document.paragraphs:
#         text.append(para.text)

#         # Check for hyperlinks in paragraph runs
#         for run in para.runs:
#             # Check for hyperlinks via the XML part (relationships)
#             if run._r.hyperlink:
#                 hyperlink = run._r.hyperlink
#                 # Extract the URL from the relationship part
#                 url = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink}t")
#                 if url:
#                     urls.append(url)
                
#     return "\n".join(text).strip(), urls


def extract_text_from_word(file_path):
    urls = []
    document = Document(file_path)
    text = "\n".join([para.text for para in document.paragraphs])
    return text.strip(), urls


def extract_text(file_path):
    if not os.path.exists(file_path):
        return None
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_extension == ".docx":
        return extract_text_from_word(file_path)
    return None